import argparse
import base64
import io
import math
import os
import sys
import zlib
from dataclasses import dataclass

CURRENT_DIR = os.path.dirname(os.path.abspath(__file__))
SRC_DIR = os.path.dirname(CURRENT_DIR)
if SRC_DIR not in sys.path:
    sys.path.insert(0, SRC_DIR)

import segno
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from reedsolo import RSCodec

from i18n import tr as i18n_tr
from i18n.core_texts import AUTO_LANGUAGE_VALUE, CLI_LANGUAGE_CHOICES, detect_lang

# ==============================================================================
#                               核心配置区域
# ==============================================================================
CHUNK_SIZE = 500
QR_ERROR = "M"
QR_WIDTH_CM = 4.0
FONT_SIZE_LABEL = 10

COLS_PER_PAGE = 4
PAGE_MARGIN = 1.0
QR_ERROR_CHOICES = ("L", "M", "Q", "H")
TARGET_ENCODING = "utf-8"
BASE64_TAG = "<<BASE64>>"
RS_TAG = "<<RS:"
# RS每次编码的最大符号数（GF(2^8)，码字总数最多255，留出冗余空间）
RS_MAX_NSZ = 255
# ==============================================================================


@dataclass(frozen=True)
class QrLayoutConfig:
    chunk_size: int = CHUNK_SIZE
    qr_error: str = QR_ERROR
    qr_width_cm: float = QR_WIDTH_CM
    font_size_label: int = FONT_SIZE_LABEL
    cols_per_page: int = COLS_PER_PAGE
    page_margin: float = PAGE_MARGIN
    enable_redundancy: bool = False
    rs_block_ratio: float = 0.05


DEFAULT_CONFIG = QrLayoutConfig()


def tr(lang: str, key: str, **kwargs) -> str:
    return i18n_tr("auto_split_qr", lang, key, **kwargs)


def _encode_input_data(raw_data: bytes) -> tuple[str, bool]:
    try:
        return raw_data.decode(TARGET_ENCODING), False
    except UnicodeDecodeError:
        return base64.b64encode(raw_data).decode("ascii"), True


def _calculate_crc32(data: str) -> str:
    """计算字符串的CRC32校验和，返回8字符的16进制字符串"""
    crc_value = zlib.crc32(data.encode("utf-8")) & 0xFFFFFFFF
    return format(crc_value, "08X")


def _reed_solomon_encode(chunks: list[str], redundancy_blocks: int) -> list[str]:
    """
    使用Reed-Solomon编码为数据块添加冗余（列编码策略）。

    将数据块视为矩阵的行，对每个字节列独立进行RS编码。
    当某些块（整行）损坏或缺失时，可利用剩余行通过RS解码恢复。
    冗余块以base64编码输出（二进制→文本，确保可安全放入QR码）。

    Args:
        chunks: 原始数据块列表（每个为字符串）
        redundancy_blocks: 冗余块数量

    Returns:
        原始数据块 + 冗余块（base64编码）
    """
    if redundancy_blocks <= 0 or not chunks:
        return chunks

    n = len(chunks)
    max_len = max(len(c.encode("utf-8")) for c in chunks)
    redundancy_rows = [bytearray() for _ in range(redundancy_blocks)]
    rs = RSCodec(redundancy_blocks)

    for col_idx in range(max_len):
        column_bytes = bytearray(n)
        for row_idx, chunk in enumerate(chunks):
            chunk_b = chunk.encode("utf-8")
            column_bytes[row_idx] = chunk_b[col_idx] if col_idx < len(chunk_b) else 0

        encoded = rs.encode(bytes(column_bytes))
        for k in range(redundancy_blocks):
            redundancy_rows[k].append(encoded[n + k])

    redundancy_strs = [base64.b64encode(bytes(r)).decode("ascii") for r in redundancy_rows]
    return chunks + redundancy_strs


def _build_qr_task(idx, total_chunks, chunk_content, is_base64, original_filename, config: QrLayoutConfig):
    """构建无CRC的QR码（原始兼容模式）"""
    if is_base64:
        if idx == 1:
            payload = f"[{idx}/{total_chunks}]<<FILENAME:{original_filename}>>{BASE64_TAG}{chunk_content}"
        else:
            payload = f"[{idx}/{total_chunks}]{chunk_content}"
    else:
        payload = f"[{idx}/{total_chunks}]{chunk_content}"
        if idx == total_chunks:
            payload += f"<<FILENAME:{original_filename}>>"

    qr = segno.make(payload, error=config.qr_error)
    img_stream = io.BytesIO()
    qr.save(img_stream, kind="png", scale=10, border=1)
    img_stream.seek(0)
    return idx, img_stream


def _build_qr_task_with_crc(idx, total_chunks, chunk_content, is_base64, original_filename,
                            config: QrLayoutConfig, original_count: int, redundancy_blocks: int):
    """构建包含CRC32校验和RS元数据的QR码"""
    rs_meta = f"{RS_TAG}{original_count};{redundancy_blocks}>>"

    # 构建数据部分（[idx/total][CRC32] 之后的内容）
    if is_base64:
        if idx == 1:
            data_part = f"<<FILENAME:{original_filename}>>{rs_meta}{BASE64_TAG}{chunk_content}"
        else:
            data_part = f"{rs_meta}{chunk_content}"
    else:
        data_part = f"{rs_meta}{chunk_content}"
        if idx == total_chunks:
            data_part += f"<<FILENAME:{original_filename}>>"

    # 方案A：对 meta + chunk_content 整体计算CRC
    crc32_checksum = _calculate_crc32(data_part)

    # 构建完整 payload
    payload = f"[{idx}/{total_chunks}][{crc32_checksum}]{data_part}"

    qr = segno.make(payload, error=config.qr_error)
    img_stream = io.BytesIO()
    qr.save(img_stream, kind="png", scale=10, border=1)
    img_stream.seek(0)
    return idx, img_stream


def setup_page(doc, config: QrLayoutConfig):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin, section.right_margin = Cm(config.page_margin), Cm(config.page_margin)
    section.top_margin, section.bottom_margin = Cm(config.page_margin), Cm(config.page_margin)


def process_file(input_path: str, lang: str = "zh", config: QrLayoutConfig | None = None):
    input_path = os.path.abspath(input_path)
    config = config or DEFAULT_CONFIG
    if not os.path.exists(input_path):
        return print(tr(lang, "missing_input", input_file=input_path))

    original_filename = os.path.basename(input_path)
    base_name, _ = os.path.splitext(input_path)
    output_doc = f"{base_name}{tr(lang, 'suffix')}.docx"

    with open(input_path, "rb") as f:
        raw_data = f.read()

    full_text, is_base64 = _encode_input_data(raw_data)
    total_chars = len(full_text)

    # 初始分块
    initial_chunks = []
    for i in range(0, total_chars, config.chunk_size):
        initial_chunks.append(full_text[i:i + config.chunk_size])

    original_chunk_count = len(initial_chunks)
    redundancy_blocks = 0

    # 处理冗余编码
    if config.enable_redundancy:
        max_allowed_redundancy = RS_MAX_NSZ - original_chunk_count
        if max_allowed_redundancy < 2:
            # 块太多无法用RS，降级为无冗余模式
            print(tr(lang, "redundancy_too_many_blocks", max_chunks=original_chunk_count))
            config = QrLayoutConfig(
                chunk_size=config.chunk_size, qr_error=config.qr_error,
                qr_width_cm=config.qr_width_cm, font_size_label=config.font_size_label,
                cols_per_page=config.cols_per_page, page_margin=config.page_margin,
                enable_redundancy=False, rs_block_ratio=config.rs_block_ratio,
            )
        else:
            requested = max(2, int(original_chunk_count * config.rs_block_ratio))
            redundancy_blocks = max(2, min(requested, max_allowed_redundancy))
            print("=" * 50)
            print(tr(lang, "data_loaded_base64" if is_base64 else "data_loaded_text"))
            print(tr(lang, "original_chars", total_chars=total_chars))
            print(tr(lang, "redundancy_enabled", redundancy_blocks=redundancy_blocks,
                       max_recoverable=redundancy_blocks))
            print(tr(lang, "generated_count",
                       total_chunks=original_chunk_count + redundancy_blocks))
            print(tr(lang, "rs_encoding_done"))
            print("=" * 50)
            all_chunks = _reed_solomon_encode(initial_chunks, redundancy_blocks)

    if not config.enable_redundancy:
        all_chunks = initial_chunks
        total_chunks = len(all_chunks)
        print("=" * 50)
        print(tr(lang, "data_loaded_base64" if is_base64 else "data_loaded_text"))
        print(tr(lang, "original_chars", total_chars=total_chars))
        print(tr(lang, "generated_count", total_chunks=total_chunks))
        print(tr(lang, "redundancy_disabled"))
        print("=" * 50)

    total_chunks = len(all_chunks)

    doc = Document()
    setup_page(doc, config)
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    total_rows = math.ceil(total_chunks / config.cols_per_page)
    current_table = doc.add_table(rows=total_rows, cols=config.cols_per_page)
    current_table.autofit = False
    progress_step = max(1, total_chunks // 100)

    for i in range(total_chunks):
        idx = i + 1
        chunk_content = all_chunks[i]

        if config.enable_redundancy:
            _, img_stream = _build_qr_task_with_crc(
                idx, total_chunks, chunk_content, is_base64, original_filename,
                config, original_chunk_count, redundancy_blocks)
        else:
            _, img_stream = _build_qr_task(
                idx, total_chunks, chunk_content, is_base64, original_filename, config)

        row_idx = i // config.cols_per_page
        col_idx = i % config.cols_per_page
        cell = current_table.cell(row_idx, col_idx)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        p_label = cell.paragraphs[0]
        p_label.alignment, p_label.paragraph_format.space_before, p_label.paragraph_format.space_after = \
            WD_ALIGN_PARAGRAPH.CENTER, Pt(0), Pt(0)
        run_label = p_label.add_run(tr(lang, "part_label", current=idx, total=total_chunks))
        run_label.font.bold, run_label.font.size = True, Pt(config.font_size_label)

        p_img = cell.add_paragraph()
        p_img.alignment, p_img.paragraph_format.space_before, p_img.paragraph_format.space_after = \
            WD_ALIGN_PARAGRAPH.CENTER, Pt(0), Pt(0)
        p_img.add_run().add_picture(img_stream, width=Cm(config.qr_width_cm))
        img_stream.close()

        if idx % progress_step == 0 or idx == total_chunks:
            print(f"__PROGRESS__::{idx}::{total_chunks}")
            print(tr(lang, "progress", current=idx, total=total_chunks))

    doc.save(output_doc)
    print(tr(lang, "saved", output_doc=output_doc))


def process_files(input_paths, lang: str = "zh", config: QrLayoutConfig | None = None):
    for input_path in input_paths:
        if not input_path:
            continue
        abs_path = os.path.abspath(input_path)
        if os.path.isdir(abs_path):
            print(tr(lang, "missing_input", input_file=abs_path))
            continue
        process_file(abs_path, lang=lang, config=config)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("-l", "--lang", choices=CLI_LANGUAGE_CHOICES, default=AUTO_LANGUAGE_VALUE)
    parser.add_argument("target", nargs="*")
    args = parser.parse_args()

    lang = detect_lang(None if args.lang == AUTO_LANGUAGE_VALUE else args.lang)
    if not args.target:
        return print(tr(lang, "missing_input", input_file="<No file provided>"))

    for target in args.target:
        target_path = os.path.abspath(target)
        if os.path.isdir(target_path):
            try:
                from core import scanner_decoder
            except ImportError:
                import scanner_decoder
            scanner_decoder.decode_folder(target_path, lang=lang)
        else:
            process_file(target_path, lang=lang)


if __name__ == "__main__":
    main()