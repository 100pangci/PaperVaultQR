import argparse
import locale
import math
import os
import shutil

import segno
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

# ==============================================================================
#                               核心配置区域
# ==============================================================================
INPUT_FILE = "split_qr.json"
OUTPUT_DOCS = {
    "zh": "冷存储.docx",
    "en": "cold_storage.docx",
}

CHUNK_SIZE = 500       # 回归最舒服的 500
QR_ERROR = "M"
QR_WIDTH_CM = 4.0
FONT_SIZE_LABEL = 10

ROWS_PER_PAGE = 6      # 4列6行 = 24个/页
COLS_PER_PAGE = 4
PAGE_MARGIN = 1.0
# ==============================================================================

TEMP_DIR = "temp_qrcodes_cache"

MESSAGES = {
    "zh": {
        "missing_input": "❌ 找不到文件 '{input_file}'",
        "data_loaded": "📄 数据加载完毕 (纯净切片模式)",
        "original_chars": "   - 原始字符数: {total_chars}",
        "generated_count": "   - 生成数量: {total_chunks} 个二维码 (约 {total_pages} 页)",
        "progress": "✅ 处理进度: {current} / {total}",
        "saved": "\n💾 成功生成打印文档: {output_doc}",
        "part_label": "第 {current} / {total} 部分",
    },
    "en": {
        "missing_input": "❌ File not found: '{input_file}'",
        "data_loaded": "📄 Data loaded (plain slicing mode)",
        "original_chars": "   - Total characters: {total_chars}",
        "generated_count": "   - Generated: {total_chunks} QR codes (about {total_pages} pages)",
        "progress": "✅ Progress: {current} / {total}",
        "saved": "\n💾 Print document generated successfully: {output_doc}",
        "part_label": "Part {current} / {total}",
    },
}


def detect_lang(explicit_lang: str | None = None) -> str:
    if explicit_lang in {"zh", "en"}:
        return explicit_lang

    for env_key in ("APP_LANG", "LANGUAGE", "LC_ALL", "LANG"):
        value = os.environ.get(env_key, "").lower()
        if value.startswith("zh") or "zh_" in value or "zh-" in value:
            return "zh"
        if value.startswith("en") or "en_" in value or "en-" in value:
            return "en"

    try:
        locale_name = locale.getlocale()[0] or ""
    except (ValueError, TypeError):
        locale_name = ""

    return "zh" if "zh" in locale_name.lower() else "en"


def tr(lang: str, key: str, **kwargs) -> str:
    template = MESSAGES[lang][key]
    return template.format(**kwargs)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Split text into QR codes and generate a Word document."
    )
    parser.add_argument(
        "-l",
        "--lang",
        choices=("auto", "zh", "en"),
        default="auto",
        help="Language for console output and generated labels.",
    )
    return parser.parse_args()


def setup_page(doc):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin, section.right_margin = Cm(PAGE_MARGIN), Cm(PAGE_MARGIN)
    section.top_margin, section.bottom_margin = Cm(PAGE_MARGIN), Cm(PAGE_MARGIN)


def main():
    args = parse_args()
    lang = detect_lang(None if args.lang == "auto" else args.lang)
    output_doc = OUTPUT_DOCS[lang]

    if not os.path.exists(INPUT_FILE):
        return print(tr(lang, "missing_input", input_file=INPUT_FILE))

    with open(INPUT_FILE, "r", encoding="utf-8") as f:
        full_text = f.read().strip()

    total_chars = len(full_text)
    total_chunks = math.ceil(total_chars / CHUNK_SIZE)
    chunks_per_page = ROWS_PER_PAGE * COLS_PER_PAGE
    total_pages = math.ceil(total_chunks / chunks_per_page) if total_chunks else 0

    print("=" * 50)
    print(tr(lang, "data_loaded"))
    print(tr(lang, "original_chars", total_chars=total_chars))
    print(tr(lang, "generated_count", total_chunks=total_chunks, total_pages=total_pages))
    print("=" * 50)

    if os.path.exists(TEMP_DIR):
        shutil.rmtree(TEMP_DIR, ignore_errors=True)
    os.makedirs(TEMP_DIR)

    doc = Document()
    setup_page(doc)
    current_table = None

    for i in range(total_chunks):
        idx = i + 1
        start = i * CHUNK_SIZE
        chunk_content = full_text[start:start + CHUNK_SIZE]

        payload = f"[{idx}/{total_chunks}]{chunk_content}"
        qr = segno.make(payload, error=QR_ERROR)
        img_path = os.path.join(TEMP_DIR, f"qr_{idx}.png")
        qr.save(img_path, scale=10, border=1)

        if i % chunks_per_page == 0:
            if i > 0:
                doc.add_page_break()
            current_table = doc.add_table(rows=ROWS_PER_PAGE, cols=COLS_PER_PAGE)
            current_table.autofit = False

        cell = current_table.cell(
            (i % chunks_per_page) // COLS_PER_PAGE,
            (i % chunks_per_page) % COLS_PER_PAGE,
        )
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        p_label = cell.paragraphs[0]
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_label.paragraph_format.space_before = Pt(0)
        p_label.paragraph_format.space_after = Pt(0)
        run_label = p_label.add_run(tr(lang, "part_label", current=idx, total=total_chunks))
        run_label.font.bold = True
        run_label.font.size = Pt(FONT_SIZE_LABEL)

        p_img = cell.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(0)
        p_img.paragraph_format.space_after = Pt(0)
        p_img.add_run().add_picture(img_path, width=Cm(QR_WIDTH_CM))

        print(tr(lang, "progress", current=idx, total=total_chunks))

    doc.save(output_doc)
    shutil.rmtree(TEMP_DIR, ignore_errors=True)
    print(tr(lang, "saved", output_doc=output_doc))


if __name__ == "__main__":
    main()
